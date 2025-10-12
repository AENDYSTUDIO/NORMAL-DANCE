#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создание документа со стандартами кодирования NormalDance
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def create_coding_standards_document():
    """Создает документ со стандартами кодирования"""

    doc = Document()

    # Настройка стилей
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(18)
    title_style.font.bold = True

    heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    # Титульная страница
    title = doc.add_paragraph('Стандарты кодирования\nNormalDance', style='CustomTitle')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('Версия 1.0', style='CustomHeading2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('Дата: Декабрь 2025', style='CustomHeading2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    # Введение
    doc.add_paragraph('1. Введение', style='CustomHeading1')

    intro_text = """
    Данный документ содержит стандарты кодирования для команды разработки NormalDance.
    Цель документа - обеспечение一致ности кода, ускорение code review и повышение качества программного обеспечения.

    Стандарты внедряются поэтапно после согласования с вице-президентом по разработке.
    Документ будет регулярно обновляться на основе обратной связи от команды.
    """

    doc.add_paragraph(intro_text.strip())

    # 2. Общие принципы
    doc.add_paragraph('2. Общие принципы кодирования', style='CustomHeading1')

    principles = [
        "• Читаемость кода превыше оптимизации производительности",
        "• DRY (Don't Repeat Yourself) - избегайте дублирования кода",
        "• KISS (Keep It Simple, Stupid) - простота превыше сложности",
        "• Явное лучше неявного (вдохновлено Python)",
        "• Код должен быть самодокументируемым",
        "• Безопасность и надежность превыше скорости разработки"
    ]

    for principle in principles:
        doc.add_paragraph(principle)

    # 3. Технологический стек и конфигурации
    doc.add_paragraph('3. Технологический стек и конфигурации', style='CustomHeading1')

    stack_text = """
    3.1 Основные технологии:
    • Frontend: React 18, Next.js 15, TypeScript 5
    • Backend: Node.js, TypeScript
    • Database: PostgreSQL (Neon), Prisma ORM
    • Testing: Jest, React Testing Library, Playwright
    • Styling: Tailwind CSS
    • Deployment: Vercel, Docker

    3.2 Конфигурации линтеров и форматтеров:
    • ESLint: next/core-web-vitals конфигурация
    • TypeScript: strict mode включен
    • Prettier: автоматическое форматирование
    """

    doc.add_paragraph(stack_text.strip())

    # 4. Стилистические стандарты
    doc.add_paragraph('4. Стилистические стандарты', style='CustomHeading1')

    style_standards = """
    4.1 Именование:
    • Переменные и функции: camelCase
    • Компоненты и классы: PascalCase
    • Константы: UPPER_SNAKE_CASE
    • Файлы: kebab-case.tsx
    • Папки: kebab-case

    4.2 Импорты:
    • React импорты отдельно от остальных
    • Группировка: React, сторонние библиотеки, внутренние модули
    • Абсолютные импорты через @/ для src/

    4.3 Компоненты React:
    • Функциональные компоненты с хуками
    • Деструктуризация пропсов
    • Ранний return для условного рендеринга
    """

    doc.add_paragraph(style_standards.strip())

    # 5. Тестирование
    doc.add_paragraph('5. Тестирование', style='CustomHeading1')

    testing_standards = """
    5.1 Обязательные тесты:
    • Unit тесты для утилит и хуков (Jest)
    • Component тесты для UI компонентов (React Testing Library)
    • Integration тесты для API endpoints
    • E2E тесты для критических пользовательских сценариев (Playwright)

    5.2 Требования к покрытию:
    • Минимум 80% покрытие для новых компонентов
    • 100% покрытие для бизнес-логики
    • Критические пути: 100% покрытие

    5.3 Соглашения по命名:
    • Тестовые файлы: *.test.ts, *.test.tsx
    • Describe блоки: описание компонента/функции
    • It блоки: конкретное поведение
    """

    doc.add_paragraph(testing_standards.strip())

    # 6. Документация
    doc.add_paragraph('6. Документация', style='CustomHeading1')

    docs_standards = """
    6.1 JSDoc комментарии:
    • Обязательны для публичных API функций
    • Сложные бизнес-логика функции
    • Пропсы компонентов с неочевидным назначением

    6.2 README файлы:
    • Обязательны для новых модулей/пакетов
    • Описание API, примеры использования
    • Инструкции по настройке и запуску

    6.3 Storybook:
    • Обязателен для переиспользуемых UI компонентов
    • Документация пропсов и вариантов использования
    """

    doc.add_paragraph(docs_standards.strip())

    # 7. Git Workflow
    doc.add_paragraph('7. Git Workflow', style='CustomHeading1')

    git_standards = """
    7.1 Branch naming:
    • feature/description-kebab-case
    • bugfix/issue-description
    • hotfix/critical-fix
    • refactor/component-improvement

    7.2 Pull Request titles:
    • [FEATURE] Добавление новой функциональности
    • [BUGFIX] Исправление ошибки в компоненте
    • [REFACTOR] Улучшение структуры кода
    • [DOCS] Обновление документации

    7.3 Commit messages:
    type(scope): description

    [optional body]

    Types: feat, fix, docs, style, refactor, test, chore
    """

    doc.add_paragraph(git_standards.strip())

    # 8. Code Review
    doc.add_paragraph('8. Code Review', style='CustomHeading1')

    review_standards = """
    8.1 Автоматизированные проверки:
    • ESLint ошибки: blocking
    • TypeScript ошибки: blocking
    • Тесты: blocking при падении
    • Prettier: auto-fix

    8.2 Ручной review:
    • Логика и алгоритмы
    • Безопасность (SQL injection, XSS)
    • Производительность
    • Читабельность и поддерживаемость

    8.3 Процесс:
    • Максимум 2-3 ревьювера
    • Обязательный approval от code owner
    • Разрешение конфликтов через обсуждение
    """

    doc.add_paragraph(review_standards.strip())

    # 9. Безопасность
    doc.add_paragraph('9. Безопасность', style='CustomHeading1')

    security_standards = """
    9.1 Валидация данных:
    • Все пользовательские входные данные
    • SQL параметры через Prisma
    • API responses sanitization

    9.2 Аутентификация и авторизация:
    • JWT tokens с expiration
    • Role-based access control
    • Secure headers (helmet.js)

    9.3 Мониторинг:
    • Rate limiting на API endpoints
    • Error logging и alerting
    • Security audit logs
    """

    doc.add_paragraph(security_standards.strip())

    # 10. Производительность
    doc.add_paragraph('10. Производительность', style='CustomHeading1')

    perf_standards = """
    10.1 Frontend:
    • Bundle size < 500KB (gzip)
    • First Contentful Paint < 2s
    • Lighthouse score > 90

    10.2 Backend:
    • API response time < 500ms
    • Database queries optimization
    • Memory leaks prevention

    10.3 Мониторинг:
    • Web Vitals tracking
    • Performance budgets
    • Automated alerts
    """

    doc.add_paragraph(perf_standards.strip())

    # Сохранение документа
    doc.save('NORMALDANCE_Coding_Standards_v1.0.docx')
    print("Документ 'NORMALDANCE_Coding_Standards_v1.0.docx' успешно создан!")

if __name__ == "__main__":
    create_coding_standards_document()
